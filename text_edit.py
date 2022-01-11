from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import re
import math
import datetime

from abbreviations import schwartz_hearst
from dateparser.search import search_dates


class xedit:
    def __init__(self, filename):
        self.Wordfile = filename
        self.year = int(datetime.datetime.today().year)
        self.YEARS = list(range(self.year, self.year - 200, -1))
        self.years_shown = []

    def getText(self):
        fullText = []
        for para in self.Wordfile.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    def bullet_points(self):
        # region Find and process bullet points
        first = True
        for p in self.Wordfile.paragraphs:
            if 'List' in p.style.name:
                for r in p.runs:
                    if first is True:
                        r.text = r.text[0].upper() + r.text[1:]
                        first = False
                    elif ';' in r.text:
                        pass
                    else:
                        r.text = r.text + ';'

            try:
                bullet_list = [p for p in self.Wordfile.paragraphs if 'list' in p.style.name]
                first_letter = bullet_list[-1].runs[0].text[0]
                bullet_list[-1].runs[-1].text = 'And ' + first_letter.lower() + bullet_list[-1].runs[0].text[1:]
            except IndexError:
                print("Couldn't capitalise first letter, skipping")
        # endregion

    def abbreviations(self):
        # region Find abbreviations
        pairs = schwartz_hearst.extract_abbreviation_definition_pairs(doc_text=self.getText())
        # endregion

        # region Add abbreviations page
        self.Wordfile.add_page_break()
        self.Wordfile.add_heading('Abbreviations: ')
        self.Wordfile.add_paragraph(pairs)
        # endregion

    def ratio_format(self, paragraph):
        # region Check if ratio is in paragraph
        for r in paragraph.runs:
            for y in re.findall(r'\b\d+\b\s+:\s+\b\d+\b', r.text):
                r.text = r.text.replace(' : ', ':')
        # endregion

    def layout(self, paragraph, font_size=11, font='Trebuchet MS', align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        # region Format each paragraph layout
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0
        paragraph_format.line_spacing = 1.15
        paragraph_format.alignment = align
        # endregion
        for r in paragraph.runs:
            # region Format font, size and colour
            r.font.size = Pt(font_size)
            r.font.name = font
            if r.font.color.rgb == RGBColor(0, 112, 192):
                pass
            if r.font.color.rgb == RGBColor(255, 255, 0):
                pass
            if r.font.color.rgb == RGBColor(0, 176, 80):
                pass
            if r.font.color.rgb == RGBColor(0, 0, 0):
                pass
            else:
                r.font.color.rgb = RGBColor(0, 0, 0)
            # endregion

    def replace_text(self, paragraph):
        # region Replace incorrect text
        if 'poor management' in paragraph.text:
            paragraph.text = paragraph.text.replace('poor management', 'inadequate management')
        if 'improper' in paragraph.text:
            paragraph.text = paragraph.text.replace('improper', 'insufficient')
        if 'inappropriate' in paragraph.text:
            paragraph.text = paragraph.text.replace('inappropriate', 'in need of improvement')
        if 'violate' in paragraph.text:
            paragraph.text = paragraph.text.replace('violate', 'non-compliance with')
        if 'Poor management' in paragraph.text:
            paragraph.text = paragraph.text.replace('Poor management', 'Inadequate management')
        if 'Improper' in paragraph.text:
            paragraph.text = paragraph.text.replace('Improper', 'Insufficient')
        if 'Inappropriate' in paragraph.text:
            paragraph.text = paragraph.text.replace('Inappropriate', 'In need of improvement')
        if 'Violate' in paragraph.text:
            paragraph.text = paragraph.text.replace('violate', 'Non-compliance with')
        # endregion

    def date_formatter(self, r):
        # region Format date from to (eg 2020/2021)
        contains_year = any(str(YEAR) in r.text for YEAR in self.YEARS)
        if contains_year is True and '/' in r.text:
            for YEAR in self.YEARS:
                if str(YEAR) in r.text:
                    if len(years_shown) < 2:
                        years_shown.append(int(YEAR))
                    else:
                        years_shown = []
                        years_shown.append(int(YEAR))
        if len(years_shown) == 2:
            if math.isclose(years_shown[1], years_shown[0], rel_tol=1):
                first_year = str(min(years_shown))
                second_year = str(max(years_shown))
                r.text = r.text.replace(first_year + '/' + second_year, first_year + '/' + second_year[-2:])
        # endregion

        # region Format dates to shorthand
        date_list = search_dates(r.text)
        if date_list != None:
            for i in range(len(date_list)):
                if ':' in date_list[i][0]:
                    pass
                else:
                    r.text = r.text.replace(str(date_list[i][0]), date_list[i][1].strftime('%b %d, %Y'))
        # endregion

        return date_list

    def percent_format(self, r):
        # region Format percentages
        if 'percent' in r.text:
            location = r.text.find('percent')
            if str(r.text[location - 1]).isnumeric() is True:
                r.text = r.text.replace('percent', '%')
            elif str(r.text[location - 2:location - 1]).isnumeric() is True:
                r.text = r.text.replace('percent', '%')
            elif str(r.text[location - 3:location - 1]).isnumeric() is True:
                r.text = r.text.replace('percent', '%')
        # endregion