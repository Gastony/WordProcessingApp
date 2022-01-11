# region Imports
import docx
from docx import *
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mainwindow import Ui_MainWindow
from page_numbers import numbering
from text_edit import xedit

import os
import re
from PyQt5 import QtCore
from num2words import num2words
# endregion

class Model:
    def __init__(self):
        self.fileName = None
        self.fileContent = ""
        self.multi_file = False
        self.run = False
        self.main = Ui_MainWindow

    # region Get and check file
    def isValid(self, fileName):
        try:
            file = open(fileName, 'r')
            file.close()
            return True
        except:
            return False

    def setFileName(self, fileName):
        if self.isValid(fileName):
            self.fileName = fileName
            self.opened = open(self.fileName, 'rb')
            self.Wordfile = Document(self.opened)
            self.fileContents = []
            for para in self.Wordfile.paragraphs:
                self.fileContents.append(para.text)

            self.plainText = '\n'.join(self.fileContents)
        else:
            self.fileContents = ""
            self.fileName = ""

    def getFileName(self):
        return self.fileName

    def getFileContents(self):
        return self.plainText
    # endregion

    def writeDoc(self, progress_bar, open_button, remove_button):
        tex_edit = xedit(self.Wordfile)

        # region Initialise progress bar
        work_text = len([p for p in self.Wordfile.paragraphs])
        work_tables = len([table for table in self.Wordfile.tables])
        total_work = work_text+work_tables
        perc_text = work_text/total_work
        perc_tables = work_tables/total_work
        if perc_text != 0:
            iter_text = 1/perc_text/100
        if perc_tables != 0:
            iter_tables = 1/perc_tables/100
        progress = 0
        # endregion

        styles_element = self.Wordfile.styles.element

        # region Set language
        rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
        lang_default = rpr_default.xpath('w:lang')[0]
        lang_default.set(docx.oxml.shared.qn('w:val'), 'en-UK')
        # endregion

        # region Set up header
        section = self.Wordfile.sections[0]
        header = section.header
        head = header.paragraphs[0]
        head.text = 'This is the header'
        # endregion

        page_tracker = 'title'

        tex_edit.bullet_points()

        word_list = []

        for p in self.Wordfile.paragraphs:
            QtCore.QCoreApplication.processEvents()

            tex_edit.ratio_format(p)
            tex_edit.layout(p)
            tex_edit.replace_text(p)

            # region Update progress bar
            if iter_text != 0:
                progress = progress+iter_text
                progress_bar.setValue(progress)
            # endregion

            for r in p.runs:
                QtCore.QCoreApplication.processEvents()

                date_list = tex_edit.date_formatter(r)

                # region numbers to words
                for s in re.findall(r'\b\d+\b', r.text):
                    if int(s) <= 10:
                        contains_year = any(str(YEAR) in r.text for YEAR in tex_edit.YEARS)
                        if date_list != None:
                            is_date = any(str(date) in r.text for date in date_list)
                        else:
                            is_date = False
                        if contains_year is True and '/' in r.text:
                            pass
                        elif r.font.underline is True:
                            pass
                        else:
                            if is_date is False:
                                r.text = r.text.replace(s, num2words(int(s)))
                # endregion

                tex_edit.percent_format(r)

        for table in self.Wordfile.tables:
            QtCore.QCoreApplication.processEvents()

            # region Update progressbar
            if iter_tables != 0:
                if progress < 100-iter_tables:
                    progress = progress + iter_tables
                    progress_bar.setValue(progress)
                else:
                    progress = 100
                    progress_bar.setValue(progress)
            # endregion

            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        tex_edit.replace_text(p)
                        tex_edit.layout(p, font_size=10, font='Trebuchet MS', align=WD_ALIGN_PARAGRAPH.LEFT)

                        for r in p.runs:
                            QtCore.QCoreApplication.processEvents()

                            # region Realign cells containing money
                            money_list = ['TZS', 'SH', 'amount', 'Amount', 'price', 'Price', '£', '$']
                            contains_money = any(str(amount) in r.text for amount in money_list)
                            if contains_money:
                                tex_edit.layout(p, font_size=10, font='Trebuchet MS', align=WD_ALIGN_PARAGRAPH.RIGHT)
                            # endregion

                            tex_edit.date_formatter(r)
                            tex_edit.percent_format(r)

        tex_edit.abbreviations()

        # region Number pages
        self.Numbering = numbering()

        if int(self.Numbering.file_length(self.fileName)) != 1:
            self.Wordfile.sections[0].footer.is_linked_to_previous = True
            self.Numbering.add_page_number(self.Wordfile.sections[0].footer.paragraphs[0].add_run())
            self.Wordfile.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            self.Wordfile.sections[0].footer.is_linked_to_previous = True
        # endregion

        # region Add footer
        footer = section.footer
        foot = footer.add_paragraph()
        footer.paragraphs[1].text = 'This is footer\t\t\t\t\t\t\t\tML/PA/PPRA/2020/21'
        # endregion

        # region Save file
        self.filename = self.fileName.split('/')[-1]
        if os.path.isdir(self.fileName.replace(self.filename, 'PROJECTS')) is False:
            os.makedirs(self.fileName.replace(self.filename, 'PROJECTS'))

        self.replacement = '/PROJECTS/'+self.filename.replace('.docx', '_001.docx')
        self.newfilename = self.fileName.replace(self.filename, self.replacement)

        if os.path.exists(self.newfilename):
            self.Wordfile.save(self.newfilename.replace('_001', '_002'))
            self.newfilename = self.newfilename.replace('_001', '_002')

        else:
            self.Wordfile.save(self.newfilename)

        self.file_to_open = self.newfilename
        self.file_to_delete = self.fileName
        self.multi_file = True
        self.run = True
        # endregion

        # region Update UI
        open_button.setEnabled(True)
        remove_button.setEnabled(True)

        if progress != 100:
            progress_bar.setValue(100)
        # endregion

    def remove(self):
        self.file_to_open = self.fileName
        self.file_to_delete = self.newfilename
        self.multi_file = True
        self.run = True

    def openDoc(self):
        self.opened.close()
        if self.run is True:
            os.startfile(self.newfilename)
        elif self.run is False:
            os.startfile(self.fileName)